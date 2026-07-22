"""AXON ingestion pipeline (MVP).

Mirrors the design doc's pipeline at demo scale:
UPLOAD -> CLASSIFY -> PARSE -> METADATA -> CHUNK -> INDEX
                 -> ENTITY EXTRACT -> RELATION DISCOVER -> KG MERGE

Documents are markdown with YAML-ish frontmatter; the P&ID arrives as the
structured output of the (simulated) vision service.
"""
from __future__ import annotations

import csv
import json
import re
from dataclasses import dataclass, field
from pathlib import Path
from collections import Counter

DATA_DIR = Path(__file__).resolve().parent.parent / "data"
UPLOADS_DIR = DATA_DIR / "uploads"

TAG_RE = re.compile(r"\b(?:[A-Z]{1,3}-\d{2,3}|SAF-\d+|SOP-\d+|BRG-\d+|MS-\d+|CE-\d+|WO-\d+|MCC-\d-\d+|HS-\d+)\b")

# Concept extraction for arbitrary uploaded documents (design pipeline stage
# ENTITY EXTRACT). Deterministic, dependency-free: surfaces the document's
# salient named concepts so the KG can be built from ANY PDF, not just ones
# that happen to mention seeded equipment tags.
_WORD = re.compile(r"[A-Za-z][A-Za-z0-9+.#-]*")
_ACRONYM = re.compile(r"^[A-Z]{2,6}s?$")                 # RAG, LCEL, PHA, PHAs
# each capital hump must be followed by a lowercase run — rejects OCR-mangled
# tokens like "revIeWS"/"aTure" while keeping RunnableLambda, MultiHeadAttention
_PASCAL = re.compile(r"^[A-Z][a-z0-9]+(?:[A-Z][a-z0-9]+)+$")   # RunnableLambda
_CAMEL = re.compile(r"^[a-z]{2,}(?:[A-Z][a-z0-9]+)+$")         # fromMessages
_CODEID = re.compile(r"^[A-Za-z][A-Za-z0-9]*\.[A-Za-z][A-Za-z0-9.]*$")  # x.from
_TITLE = re.compile(r"^[A-Z][a-z]{2,}$")                 # plain TitleCase word

# Words that are TitleCase but not concepts: common English + book scaffolding.
_CONCEPT_STOP = {
    "the", "and", "for", "with", "this", "that", "these", "those", "from",
    "when", "where", "which", "what", "how", "why", "you", "your", "our",
    "they", "their", "his", "her", "but", "such", "each", "any", "all", "some",
    "more", "most", "other", "one", "two", "use", "used", "using", "via", "per",
    "also", "here", "there", "see", "new", "like", "etc", "get", "set", "let",
    "however", "every", "true", "false", "real", "build", "building", "code",
    "project", "practical", "yourself", "larger", "worked", "going", "deeper",
    "deep", "key", "common", "takeaways", "pitfall", "pitfalls", "practice",
    "exercise", "interview", "questions", "question", "design", "overview",
    "summary", "introduction", "conclusion", "contents", "appendix", "guide",
    "book", "plan", "basics", "basic", "advanced", "chapter", "section", "part",
    "step", "page", "figure", "table", "example", "note", "first", "second",
    "next", "before", "after", "january", "february", "march", "april", "may",
    "june", "july", "august", "september", "october", "november", "december",
    "because", "therefore", "thus", "hence", "while", "since", "although",
    "whereas", "moreover", "furthermore", "additionally", "finally", "given",
    "unlike", "instead", "meanwhile", "however",
}


def _is_concept_token(w: str) -> bool:
    """High-precision structural test: technical concept names look like
    acronyms, CamelCase, or code identifiers — not ordinary TitleCase words."""
    lw = w.lower()
    if lw.startswith("www.") or lw.endswith((".com", ".org", ".net", ".io")):
        return False                                    # URLs are not concepts
    if _ACRONYM.match(w) or _PASCAL.match(w) or _CAMEL.match(w):
        return True
    if _CODEID.match(w) and len(w) >= 5:
        return True
    return False


def _extract_concepts(texts: list[str], top_k: int = 12) -> list[str]:
    """Return the document's salient concepts. Precision-first: structural
    terms (CamelCase / acronyms / code ids) plus TitleCase bigrams whose words
    are not scaffolding. Plain single TitleCase words are dropped as noise."""
    from collections import Counter
    df: Counter = Counter()  # concept -> number of chunks it appears in
    for t in texts:
        toks = _WORD.findall(t)
        found: set[str] = set()
        for w in toks:
            if len(w) >= 3 and _is_concept_token(w):
                found.add(w)
        for i in range(len(toks) - 1):                    # TitleCase bigrams
            a, b = toks[i], toks[i + 1]
            if (_TITLE.match(a) and _TITLE.match(b)
                    and a.lower() not in _CONCEPT_STOP
                    and b.lower() not in _CONCEPT_STOP):
                found.add(f"{a} {b}")
        for c in found:
            df[c] += 1

    min_df = 2 if len(texts) >= 3 else 1                  # short PDFs: allow df=1
    cands = {c: f for c, f in df.items() if f >= min_df}
    in_bigram = {p for c in cands if " " in c for p in c.split()}
    scored = [(c, f + (1.0 if " " in c else 0.0))         # prefer phrases
              for c, f in cands.items()
              if not (" " not in c and c in in_bigram)]    # drop redundant unigrams
    scored.sort(key=lambda x: x[1], reverse=True)
    return [c for c, _ in scored[:top_k]]




def _extract_keywords(text: str, top_k: int = 12) -> list[str]:

    words = re.findall(r"[A-Za-z][A-Za-z0-9\-]+", text.lower())

    stop = {
        "the","and","for","with","this","that","from","into",
        "their","using","used","into","have","has","were",
        "been","being","will","shall","would","should",
        "about","than","also","such","there","these","those",
        "your","they","them","then","which","where","when",
        "what","does","each","other","more","most","very"
    }

    words = [
        w
        for w in words
        if len(w) > 2 and w not in stop
    ]

    return [
        w
        for w, _ in Counter(words).most_common(top_k)
    ]


@dataclass
class Chunk:

    chunk_id: str

    doc_no: str

    doc_title: str

    revision: str

    page: int

    section: str

    subsection: str = ""

    chunk_type: str = "text"

    text: str = ""

    summary: str = ""

    entities: list[str] = field(default_factory=list)

    concepts: list[str] = field(default_factory=list)

    keywords: list[str] = field(default_factory=list)

    token_count: int = 0

@dataclass
class Corpus:
    chunks: list[Chunk]
    docs: dict            # doc_no -> metadata
    pid: dict             # parsed P&ID topology
    maintenance: list[dict]
    spares: list[dict]
    sensors: list[dict]
    assets: dict = field(default_factory=dict)  # generated Asset360 register


def _parse_frontmatter(raw: str) -> tuple[dict, str]:
    meta: dict = {}
    body = raw
    if raw.startswith("---"):
        _, fm, body = raw.split("---", 2)
        for line in fm.strip().splitlines():
            if ":" not in line:
                continue
            k, v = line.split(":", 1)
            v = v.strip()
            if v.startswith("[") and v.endswith("]"):
                meta[k.strip()] = [x.strip() for x in v[1:-1].split(",") if x.strip()]
            else:
                meta[k.strip()] = v
    return meta, body.strip()


def _summarize_chunk(text: str) -> str:
    """
    Lightweight summary.

    We intentionally avoid using an LLM during ingestion.
    """

    paragraphs = [
        p.strip()
        for p in text.split("\n")
        if p.strip()
    ]

    if not paragraphs:
        return ""

    summary = paragraphs[0]

    if len(summary) > 250:
        summary = summary[:250] + "..."

    return summary


def _extract_pdf_text(page) -> str:
    """Extract PDF text with tighter character spacing.

    Some uploaded papers contain compact glyph positioning. pdfplumber's
    default tolerance can glue words together ("TheEraof1-bitLLMs"), which
    damages keyword search and makes real prose look like outline noise.
    """
    text = page.extract_text(x_tolerance=1, y_tolerance=3) or ""
    return text.strip()


def _chunk_markdown(doc_no: str, title: str, revision: str, body: str) -> list[Chunk]:
    """
    Structure-aware markdown chunking.

    One chunk per markdown section (## Heading), enriched with metadata.
    """

    chunks = []

    sections = re.split(r"\n(?=## )", body)

    for i, sec in enumerate(sections):

        sec = sec.strip()

        if not sec:
            continue

        first_line = sec.splitlines()[0].lstrip("# ").strip()

        section_name = first_line if sec.startswith("#") else "Preamble"

        chunk = Chunk(
            chunk_id=f"{doc_no}::{i}",
            doc_no=doc_no,
            doc_title=title,
            revision=revision,

            page=1,                          # markdown has no page numbers
            section=section_name,
            subsection="",
            chunk_type="markdown",

            text=sec,

            summary=_summarize_chunk(sec),

            entities=sorted(set(TAG_RE.findall(sec))),

            concepts=[],

            keywords=_extract_keywords(sec),

            token_count=len(sec.split())
        )

        chunks.append(chunk)

    return chunks

def _strip_running_headers(page_texts: list[tuple[int, str]]
                           ) -> list[tuple[int, str]]:
    """Remove running page headers from extracted PDF pages.

    A book's header ("LangChain Complete Knowledge Base | For AI Engineer
    Roles") is repeated as the first line of nearly every page. Left in the
    chunk text it is poison twice over: the generator reads it as content
    (producing 'LangChain is a comprehensive knowledge base for AI
    engineers' — confusing the library with the DOCUMENT TITLE), and the
    claim validator then finds lexical support for that wrong claim in
    every chunk. Detection: a first line (or its 4-word prefix, for headers
    with varying chapter suffixes) that repeats on >=30% of pages is a
    running header, not content."""
    if len(page_texts) < 4:
        return page_texts
    firsts = [t.split("\n", 1)[0].strip() for _, t in page_texts if t.strip()]
    n = len(firsts) or 1
    threshold = max(3, int(0.3 * n))
    exact = Counter(firsts)
    prefix = Counter(
        " ".join(f.split()[:4]) for f in firsts if len(f.split()) >= 4)

    def clean(text: str) -> str:
        head, sep, rest = text.partition("\n")
        h = head.strip()
        if not h:
            return text
        if exact[h] >= threshold:
            return rest
        words = h.split()
        if len(words) >= 4:
            p = " ".join(words[:4])
            if prefix[p] >= threshold:
                remainder = " ".join(words[4:])
                return (remainder + sep + rest) if remainder else rest
        return text

    return [(pno, clean(t)) for pno, t in page_texts]


def _chunk_pdf(pdf_path: Path) -> tuple[dict, list[Chunk]]:
    """
    Ingest an uploaded PDF.

    Improvements:
    - Paragraph-aware chunking
    - Rich metadata
    - Keyword extraction
    - Lightweight summaries
    - Running-header removal
    """

    import pdfplumber

    doc_no = re.sub(r"[^A-Za-z0-9_-]+", "-", pdf_path.stem)[:40]

    chunks: list[Chunk] = []

    mentions: set[str] = set()

    with pdfplumber.open(pdf_path) as pdf:
        page_texts = [
            (pno, text)
            for pno, page in enumerate(pdf.pages, start=1)
            if (text := _extract_pdf_text(page))
        ]

    page_texts = _strip_running_headers(page_texts)

    if True:  # preserve original indentation of the chunking loop below

        for pno, text in page_texts:

            if not text:
                continue

            paragraphs = [
                p.strip()
                for p in text.split("\n\n")
                if p.strip()
            ]

            current = ""

            section = f"Page {pno}"

            chunk_index = 0

            for para in paragraphs:

                # Detect headings
                if len(para) < 80 and para == para.title():

                    section = para

                if len(current) + len(para) > 1200 and current:

                    ents = sorted(set(TAG_RE.findall(current)))

                    mentions.update(ents)

                    chunks.append(
                        Chunk(
                            chunk_id=f"{doc_no}::p{pno}.{chunk_index}",
                            doc_no=doc_no,
                            doc_title=pdf_path.name,
                            revision="uploaded",

                            page=pno,

                            section=section,

                            subsection="",

                            chunk_type="text",

                            text=current,

                            summary=_summarize_chunk(current),

                            entities=ents,

                            concepts=[],

                            keywords=_extract_keywords(current),

                            token_count=len(current.split()),
                        )
                    )

                    chunk_index += 1

                    current = para

                else:

                    current += "\n\n" + para if current else para

            if current:

                ents = sorted(set(TAG_RE.findall(current)))

                mentions.update(ents)

                chunks.append(
                    Chunk(
                        chunk_id=f"{doc_no}::p{pno}.{chunk_index}",
                        doc_no=doc_no,
                        doc_title=pdf_path.name,
                        revision="uploaded",

                        page=pno,

                        section=section,

                        subsection="",

                        chunk_type="text",

                        text=current,

                        summary=_summarize_chunk(current),

                        entities=ents,

                        concepts=[],

                        keywords=_extract_keywords(current),

                        token_count=len(current.split()),
                    )
                )

    concepts = _extract_concepts(
        [c.text for c in chunks]
    ) if chunks else []

    concept_res = [
        (
            c,
            re.compile(r"\b" + re.escape(c) + r"\b", re.I)
        )
        for c in concepts
    ]

    for ch in chunks:

        present = [
            c
            for c, rx in concept_res
            if rx.search(ch.text)
        ]

        if present:

            ch.concepts = present

            ch.entities = sorted(
                set(ch.entities) | set(present)
            )

    meta = {
        "title": pdf_path.name,
        "type": "Uploaded",
        "revision": "uploaded",
        "mentions": sorted(mentions),
        "concepts": concepts,
    }

    return {doc_no: meta}, chunks

# ---------------------------------------------------------------------------
# Multi-format upload handlers
#
# Each handler: Path -> (meta_map, chunks). New formats register here; the
# upload endpoint and load_corpus() both dispatch through UPLOAD_HANDLERS,
# so adding a format is one extractor + one registry line.
# ---------------------------------------------------------------------------

def _pack_text_chunks(doc_no: str, title: str, text: str,
                      doc_type: str) -> tuple[dict, list[Chunk]]:
    """Generic text -> chunks: paragraph packing (~1200 chars), keywords,
    corpus-wide concept extraction, entity tags — the same enrichment the
    PDF path gets, for any format that can yield plain text."""
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    chunks: list[Chunk] = []
    mentions: set[str] = set()
    current = ""
    part = 1

    def flush():
        nonlocal current, part
        if not current:
            return
        ents = sorted(set(TAG_RE.findall(current)))
        mentions.update(ents)
        chunks.append(Chunk(
            chunk_id=f"{doc_no}::s{part}",
            doc_no=doc_no, doc_title=title, revision="uploaded",
            page=part, section=f"Part {part}", chunk_type="text",
            text=current, summary=_summarize_chunk(current),
            entities=ents, concepts=[],
            keywords=_extract_keywords(current),
            token_count=len(current.split()),
        ))
        part += 1
        current = ""

    for para in paragraphs:
        if len(current) + len(para) > 1200 and current:
            flush()
        current += ("\n\n" + para) if current else para
    flush()

    concepts = _extract_concepts([c.text for c in chunks]) if chunks else []
    concept_res = [(c, re.compile(r"\b" + re.escape(c) + r"\b", re.I))
                   for c in concepts]
    for ch in chunks:
        present = [c for c, rx in concept_res if rx.search(ch.text)]
        if present:
            ch.concepts = present
            ch.entities = sorted(set(ch.entities) | set(present))

    meta = {"title": title, "type": doc_type, "revision": "uploaded",
            "mentions": sorted(mentions), "concepts": concepts}
    return {doc_no: meta}, chunks


def _doc_no_for(path: Path) -> str:
    return re.sub(r"[^A-Za-z0-9_-]+", "-", path.stem)[:40]


def _extract_html_text(path: Path) -> str:
    from html.parser import HTMLParser

    class _Stripper(HTMLParser):
        SKIP = {"script", "style", "noscript"}
        BLOCK = {"p", "div", "br", "li", "tr", "h1", "h2", "h3", "h4",
                 "h5", "h6", "section", "article", "table"}

        def __init__(self):
            super().__init__()
            self.parts: list[str] = []
            self._skip = 0

        def handle_starttag(self, tag, attrs):
            if tag in self.SKIP:
                self._skip += 1
            elif tag in self.BLOCK:
                self.parts.append("\n\n")

        def handle_endtag(self, tag):
            if tag in self.SKIP and self._skip:
                self._skip -= 1

        def handle_data(self, data):
            if not self._skip:
                self.parts.append(data)

    stripper = _Stripper()
    stripper.feed(path.read_text(encoding="utf-8", errors="replace"))
    return "".join(stripper.parts)


def _ingest_html(path: Path) -> tuple[dict, list[Chunk]]:
    return _pack_text_chunks(_doc_no_for(path), path.name,
                             _extract_html_text(path), "Uploaded HTML")


def _ingest_docx(path: Path) -> tuple[dict, list[Chunk]]:
    import docx  # python-docx
    d = docx.Document(str(path))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return _pack_text_chunks(_doc_no_for(path), path.name,
                             "\n\n".join(parts), "Uploaded Word")


def _ingest_xlsx(path: Path, max_rows: int = 500) -> tuple[dict, list[Chunk]]:
    """Excel workbooks: every sheet, first row as header, rows rendered as
    'header: value' lines (same searchable shape as the CSV handler), with
    the sheet name kept as context."""
    import openpyxl
    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        rows = ws.iter_rows(values_only=True)
        header = next(rows, None)
        if header is None:
            continue
        header = [str(h) if h is not None else f"col{i+1}"
                  for i, h in enumerate(header)]
        lines = [f"Sheet: {ws.title}", f"Columns: {', '.join(header)}"]
        for n, row in enumerate(rows):
            if n >= max_rows:
                lines.append(f"... ({ws.title}: further rows truncated)")
                break
            cells = [f"{h}: {v}" for h, v in zip(header, row)
                     if v is not None and str(v).strip()]
            if cells:
                lines.append("; ".join(cells))
        parts.append("\n".join(lines))
    wb.close()
    return _pack_text_chunks(_doc_no_for(path), path.name,
                             "\n\n".join(parts), "Uploaded Excel")


def _ingest_csv(path: Path, max_rows: int = 500) -> tuple[dict, list[Chunk]]:
    delimiter = "\t" if path.suffix.lower() == ".tsv" else ","
    with open(path, newline="", encoding="utf-8", errors="replace") as f:
        rows = list(csv.reader(f, delimiter=delimiter))
    if not rows:
        return _pack_text_chunks(_doc_no_for(path), path.name, "", "Uploaded CSV")
    header = rows[0]
    lines = []
    for row in rows[1:max_rows + 1]:
        # header: value pairs — searchable AND readable for the LLM
        lines.append("; ".join(f"{h}: {v}" for h, v in zip(header, row) if v))
    text = f"Columns: {', '.join(header)}\n\n" + "\n".join(lines)
    return _pack_text_chunks(_doc_no_for(path), path.name, text, "Uploaded CSV")


def _ingest_json_file(path: Path, limit: int = 60_000) -> tuple[dict, list[Chunk]]:
    raw = json.loads(path.read_text(encoding="utf-8", errors="replace"))
    text = json.dumps(raw, indent=2, ensure_ascii=False)[:limit]
    return _pack_text_chunks(_doc_no_for(path), path.name, text, "Uploaded JSON")


def _ingest_image(path: Path) -> tuple[dict, list[Chunk]]:
    """Images: OCR text when an engine is available, image metadata always.
    Degrades honestly — without tesseract the image is still indexed and
    listed, and its chunk says exactly how to enable text extraction."""
    meta_lines = [f"Image file: {path.name}"]
    try:
        from PIL import Image
        with Image.open(path) as im:
            meta_lines.append(
                f"Format: {im.format}, size: {im.width}x{im.height}")
    except Exception:
        pass
    ocr_text = ""
    try:
        import shutil
        import pytesseract
        from PIL import Image
        # Servers launched from IDEs/launchers often have a minimal PATH
        # without Homebrew; locate the tesseract binary explicitly so OCR
        # works regardless of how the process was started.
        if not shutil.which("tesseract"):
            for cand in ("/opt/homebrew/bin/tesseract",
                         "/usr/local/bin/tesseract"):
                if Path(cand).exists():
                    pytesseract.pytesseract.tesseract_cmd = cand
                    break
        with Image.open(path) as im:
            ocr_text = pytesseract.image_to_string(im).strip()
    except ImportError:
        meta_lines.append(
            "OCR not available — install tesseract (brew install tesseract) "
            "and pytesseract (pip install pytesseract) to index text inside "
            "images.")
    except Exception as exc:
        meta_lines.append(f"OCR failed: {exc}")
    # When OCR succeeded, the chunk text is the DOCUMENT CONTENT ONLY.
    # Ingestion metadata ("Image file: x.png", "Format: PNG, size: ...",
    # "Extracted text (OCR):") polluted retrieval and the generator quoted
    # it as if it were what the document says ("The resume mentions an
    # image file named person_resume.png..."). The metadata-only body
    # remains for images with no extractable text — there it IS the only
    # honest content.
    if ocr_text:
        return _pack_text_chunks(_doc_no_for(path), path.name, ocr_text,
                                 "Uploaded Image")
    body = "\n".join(meta_lines)
    return _pack_text_chunks(_doc_no_for(path), path.name, body,
                             "Uploaded Image")


def _ingest_markdown_file(path: Path) -> tuple[dict, list[Chunk]]:
    meta, body = _parse_frontmatter(path.read_text())
    doc_no = meta.get("doc_no", path.stem)
    meta.setdefault("title", path.stem)
    return {doc_no: meta}, _chunk_markdown(
        doc_no, meta.get("title", path.stem), meta.get("revision", "?"), body)


UPLOAD_HANDLERS: dict[str, object] = {
    ".pdf": _chunk_pdf,
    ".md": _ingest_markdown_file,
    ".txt": _ingest_markdown_file,
    ".html": _ingest_html,
    ".htm": _ingest_html,
    ".docx": _ingest_docx,
    ".csv": _ingest_csv,
    ".tsv": _ingest_csv,
    ".xlsx": _ingest_xlsx,
    ".json": _ingest_json_file,
    ".png": _ingest_image,
    ".jpg": _ingest_image,
    ".jpeg": _ingest_image,
    ".webp": _ingest_image,
}


def supported_upload_extensions() -> tuple[str, ...]:
    return tuple(sorted(UPLOAD_HANDLERS))


def load_corpus() -> Corpus:
    chunks: list[Chunk] = []
    docs: dict = {}

    for md in sorted(list((DATA_DIR / "sops").glob("*.md")) + list((DATA_DIR / "manuals").glob("*.md"))):
        meta, body = _parse_frontmatter(md.read_text())
        doc_no = meta.get("doc_no", md.stem)
        docs[doc_no] = meta
        chunks.extend(_chunk_markdown(doc_no, meta.get("title", md.stem), meta.get("revision", "?"), body))

    # User-uploaded documents — any format with a registered handler
    # (PDF, markdown, HTML, Word, CSV, JSON, images, ...).
    UPLOADS_DIR.mkdir(exist_ok=True)
    for up in sorted(UPLOADS_DIR.iterdir()):
        handler = UPLOAD_HANDLERS.get(up.suffix.lower())
        if handler is None:
            if up.is_file() and not up.name.startswith("."):
                print(f"ingest: skipping {up.name} — no handler for "
                      f"'{up.suffix}' (supported: "
                      f"{', '.join(supported_upload_extensions())})")
            continue
        try:
            meta_map, up_chunks = handler(up)
        except Exception as exc:
            print(f"ingest: failed to parse {up.name}: {exc}")
            continue
        for m in meta_map.values():
            m["uploaded"] = True
            m["source_file"] = up.name
        docs.update(meta_map)
        chunks.extend(up_chunks)
        if up.suffix.lower() == ".pdf":
            # preserve the original PDF ordering behaviour exactly
            chunks.sort(key=lambda c: (c.doc_no, c.page, c.chunk_id))

    # Generated Asset360 artifacts are ordinary, source-cited markdown
    # documents. Loading them into the same corpus makes chat retrieval use
    # manuals, reports and timelines alongside the original upload.
    assets: dict = {}
    assets_root = DATA_DIR / "assets"
    if assets_root.exists():
        for folder in sorted(p for p in assets_root.iterdir() if p.is_dir()):
            metadata_path = folder / "metadata.json"
            try:
                metadata = json.loads(metadata_path.read_text(encoding="utf-8"))
            except (OSError, json.JSONDecodeError):
                metadata = {}
            asset_id = str(metadata.get("asset_id") or folder.name).upper()
            if metadata:
                assets[asset_id] = metadata
            for md in sorted(folder.glob("*.md")):
                try:
                    meta, body = _parse_frontmatter(md.read_text(encoding="utf-8"))
                except OSError:
                    continue
                doc_no = meta.get("doc_no", f"ASSET-{asset_id}-{md.stem}")
                meta["generated"] = True
                meta["asset_id"] = asset_id
                docs[doc_no] = meta
                chunks.extend(_chunk_markdown(doc_no, meta.get("title", md.stem),
                                              meta.get("revision", "generated"), body))

    pid = json.loads((DATA_DIR / "pid" / "pid_area1.json").read_text())

    with open(DATA_DIR / "maintenance_log.csv") as f:
        maintenance = list(csv.DictReader(f))

    # Overlay the structured maintenance_events store (history extracted from
    # uploaded documents) on top of the legacy CSV, keyed by work order. Events
    # carry richer fields (source_document, page_number, confidence, engineer,
    # ...) so this both enriches their CSV mirror rows and adds any that are not
    # yet mirrored — without ever double counting.
    from history_repository import (MaintenanceEventRepository,
                                    to_maintenance_row)
    _repo = MaintenanceEventRepository(
        DATA_DIR / "maintenance_events.json", DATA_DIR / "maintenance_log.csv")
    _by_wo = {row.get("wo_number"): i for i, row in enumerate(maintenance)}
    for _ev in _repo.all():
        _row = to_maintenance_row(_ev)
        _wo = _row["wo_number"]
        if _wo in _by_wo:
            maintenance[_by_wo[_wo]] = _row
        else:
            _by_wo[_wo] = len(maintenance)
            maintenance.append(_row)

    with open(DATA_DIR / "spares.csv") as f:
        spares = list(csv.DictReader(f))
    with open(DATA_DIR / "sensors_p101.csv") as f:
        sensors = [
            {"timestamp": r["timestamp"],
             "vibration_mm_s": float(r["vibration_mm_s"]),
             "bearing_temp_c": float(r["bearing_temp_c"]),
             "discharge_pressure_bar": float(r["discharge_pressure_bar"])}
            for r in csv.DictReader(f)
        ]

    return Corpus(chunks=chunks, docs=docs, pid=pid, maintenance=maintenance,
                  spares=spares, sensors=sensors, assets=assets)
